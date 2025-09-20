from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import streamlit as st
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import tempfile
import os
import pandas as pd
import numpy as np
from datetime import datetime

def export_report(df, keywords=None, lang="fr", include_emotions=True, include_socio=True, include_context=True):
    """
    Génère un rapport Word contenant :
      - Un tableau récapitulatif des avis, sentiments, scores et analyses avancées
      - Les analyses des émotions et socio-émotionnelles
      - La liste des thèmes récurrents (mots-clés TF-IDF)
      - Des visualisations (nuage de mots, distribution des sentiments)
    
    Parameters:
      - df: DataFrame contenant les avis et analyses
      - keywords: Liste des mots-clés TF-IDF (thèmes récurrents)
      - lang: Langue de l'analyse ('fr' ou 'en')
      - include_emotions: Inclure l'analyse des émotions si disponible
      - include_socio: Inclure l'analyse socio-émotionnelle si disponible
      - include_context: Inclure l'analyse de contexte si disponible
    
    Returns:
      - BytesIO: Fichier Word en mémoire pour téléchargement
    """
    doc = Document()
    
    # Titre et en-tête
    title = doc.add_heading("Rapport d'Analyse des Avis Patients", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajout de la date de génération
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(f"Généré le: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    # Introduction
    doc.add_paragraph(
        "Ce rapport présente une analyse détaillée des avis de patients utilisant diverses "
        "techniques de traitement automatique du langage naturel (NLP). Il inclut l'analyse "
        "de sentiment, l'extraction des thèmes récurrents, et des visualisations pertinentes."
    )
    
    # Ajout d'un résumé
    doc.add_heading("Résumé de l'analyse", level=1)
    sentiment_counts = df["Sentiment"].value_counts() if "Sentiment" in df.columns else {}
    
    p = doc.add_paragraph()
    p.add_run(f"Nombre total d'avis analysés: {len(df)}\n")
    
    if "Sentiment" in df.columns:
        p.add_run(f"Avis positifs: {sentiment_counts.get('Positif', 0)} ")
        p.add_run(f"({sentiment_counts.get('Positif', 0)/len(df)*100:.1f}%)\n")
        p.add_run(f"Avis neutres: {sentiment_counts.get('Neutre', 0)} ")
        p.add_run(f"({sentiment_counts.get('Neutre', 0)/len(df)*100:.1f}%)\n")
        p.add_run(f"Avis négatifs: {sentiment_counts.get('Négatif', 0)} ")
        p.add_run(f"({sentiment_counts.get('Négatif', 0)/len(df)*100:.1f}%)\n")
    
    # Ajouter Score moyen de sentiment
    if "Score" in df.columns:
        p.add_run(f"\nScore moyen de sentiment: {df['Score'].mean():.2f} (entre -1 et 1)\n")
    
    # Ajout des thèmes récurrents
    if keywords:
        doc.add_heading("Thèmes récurrents (TF-IDF)", level=1)
        p = doc.add_paragraph("Les mots-clés suivants apparaissent fréquemment dans les avis et sont caractéristiques du corpus: ")
        keyword_run = p.add_run(", ".join(keywords))
        keyword_run.italic = True
        keyword_run.font.color.rgb = RGBColor(0, 0, 128)  # Bleu foncé
    
    # Analyse des émotions si disponible
    emotion_cols = [col for col in df.columns if col.startswith("Emotion_")]
    if include_emotions and emotion_cols:
        doc.add_heading("Analyse des émotions", level=1)
        p = doc.add_paragraph("Cette section présente les émotions dominantes détectées dans les avis. ")
        p.add_run("Les scores représentent l'intensité moyenne de chaque émotion sur une échelle de 0 à 1.")
        
        emotion_means = df[emotion_cols].mean().sort_values(ascending=False)
        
        # Tableau des émotions
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Émotion"
        hdr_cells[1].text = "Score moyen"
        
        for emotion, score in emotion_means.items():
            emotion_name = emotion.replace("Emotion_", "").replace("_", " ")
            row_cells = table.add_row().cells
            row_cells[0].text = emotion_name
            row_cells[1].text = f"{score:.3f}"
        
        # Visualisation des émotions
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            plt.figure(figsize=(10, 6))
            emotion_means.plot(kind="bar")
            plt.title("Émotions moyennes dans les avis")
            plt.ylabel("Score moyen")
            plt.xticks(rotation=45, ha="right")
            plt.tight_layout()
            plt.savefig(tmp.name, format='png')
            plt.close()
            
            # Ajouter l'image au document
            doc.add_picture(tmp.name, width=Inches(6))
            doc.add_paragraph("Figure: Distribution des émotions détectées dans les avis.")
        
        # Supprimer le fichier temporaire
        try:
            os.unlink(tmp.name)
        except:
            pass
    
    # Analyse socio-émotionnelle si disponible
    socio_cols = [col for col in df.columns if col.startswith("Socio_") or col.startswith("Dim_")]
    if include_socio and socio_cols:
        doc.add_heading("Analyse socio-émotionnelle", level=1)
        p = doc.add_paragraph(
            "Cette section présente l'analyse des dimensions socio-émotionnelles détectées dans les avis, "
            "basée sur des modèles psychologiques comme la théorie de l'auto-détermination et le modèle PERMA."
        )
        
        dim_cols = [col for col in socio_cols if col.startswith("Dim_")]
        if dim_cols:
            # Tableau des dimensions principales
            doc.add_heading("Dimensions principales", level=2)
            dim_means = df[dim_cols].mean().sort_values(ascending=False)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Dimension"
            hdr_cells[1].text = "Score moyen"
            
            for dim, score in dim_means.items():
                dim_name = dim.replace("Dim_", "").replace("_", " ")
                row_cells = table.add_row().cells
                row_cells[0].text = dim_name
                row_cells[1].text = f"{score:.3f}"
        
        # Catégories socio-émotionnelles individuelles
        cat_cols = [col for col in socio_cols if col.startswith("Socio_")]
        if cat_cols:
            doc.add_heading("Catégories socio-émotionnelles", level=2)
            cat_means = df[cat_cols].mean().sort_values(ascending=False)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Catégorie"
            hdr_cells[1].text = "Score moyen"
            
            for cat, score in cat_means.items():
                cat_name = cat.replace("Socio_", "").replace("_", " ")
                row_cells = table.add_row().cells
                row_cells[0].text = cat_name
                row_cells[1].text = f"{score:.3f}"
    
    # Analyse de contexte si disponible
    context_cols = [col for col in df.columns if col.startswith("Context_")]
    if include_context and context_cols:
        doc.add_heading("Analyse de contexte", level=1)
        p = doc.add_paragraph(
            "Cette section présente l'analyse contextuelle des avis, incluant les entités nommées "
            "et les structures syntaxiques détectées."
        )
        
        # Tableau des entités nommées si disponible
        entity_cols = [col for col in context_cols if "entity" in col.lower()]
        if entity_cols:
            doc.add_heading("Entités nommées fréquentes", level=2)
            
            # Aggregation des entités (si stockées en format texte/liste)
            all_entities = []
            entity_types = set()
            
            for col in entity_cols:
                for cell in df[col].dropna():
                    if isinstance(cell, str):
                        # Essayer de convertir la chaîne en liste
                        try:
                            import ast
                            entities = ast.literal_eval(cell)
                            if isinstance(entities, list):
                                all_entities.extend(entities)
                                entity_types.update([e[1] for e in entities if isinstance(e, tuple) and len(e) > 1])
                        except:
                            pass
            
            if all_entities:
                # Tableau des types d'entités et leur fréquence
                from collections import Counter
                entity_counts = Counter([e[1] for e in all_entities if isinstance(e, tuple) and len(e) > 1])
                
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Type d'entité"
                hdr_cells[1].text = "Fréquence"
                
                for entity_type, count in entity_counts.most_common(10):
                    row_cells = table.add_row().cells
                    row_cells[0].text = entity_type
                    row_cells[1].text = str(count)
    
    # Ajout du nuage de mots
    doc.add_heading("Nuage de mots", level=1)
    
    # Génération du nuage de mots
    corpus = " ".join(df["review"].tolist())
    
    # Définir les stop words selon la langue
    if lang == "fr":
        try:
            import nltk
            try:
                nltk.data.find('corpora/stopwords')
            except LookupError:
                nltk.download('stopwords', quiet=True)
            from nltk.corpus import stopwords
            stop_words = list(stopwords.words('french'))
        except:
            stop_words = ["le", "la", "les", "un", "une", "des", "et", "est", "en", "que", "qui", 
                        "pour", "dans", "ce", "cette", "ces", "il", "elle", "ils", "elles", 
                        "nous", "vous", "je", "tu", "on", "son", "sa", "ses", "mon", "ma", "mes"]
    else:
        try:
            import nltk
            try:
                nltk.data.find('corpora/stopwords')
            except LookupError:
                nltk.download('stopwords', quiet=True)
            from nltk.corpus import stopwords
            stop_words = list(stopwords.words('english'))
        except:
            stop_words = "english"
    
    wordcloud = WordCloud(width=800, height=400, 
                         background_color="white", 
                         stopwords=stop_words,
                         max_words=100).generate(corpus)
    
    # Sauvegarder le nuage de mots en tant qu'image temporaire
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation="bilinear")
        plt.axis("off")
        plt.tight_layout(pad=0)
        plt.savefig(tmp.name, format='png')
        plt.close()
        
        # Ajouter l'image au document
        doc.add_picture(tmp.name, width=Inches(6))
        doc.add_paragraph("Figure: Nuage de mots généré à partir de l'ensemble des avis.")
    
    # Supprimer le fichier temporaire
    try:
        os.unlink(tmp.name)
    except:
        pass
    
    # Distribution des sentiments
    if "Sentiment" in df.columns:
        doc.add_heading("Distribution des sentiments", level=1)
        
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            plt.figure(figsize=(10, 6))
            # Utiliser les couleurs appropriées pour les sentiments
            colors = {"Positif": "green", "Neutre": "gray", "Négatif": "red"}
            sentiment_counts.plot(kind='bar', color=[colors.get(x, 'blue') for x in sentiment_counts.index])
            plt.title("Distribution des sentiments")
            plt.ylabel("Nombre d'avis")
            plt.tight_layout()
            plt.savefig(tmp.name, format='png')
            plt.close()
            
            # Ajouter l'image au document
            doc.add_picture(tmp.name, width=Inches(6))
            doc.add_paragraph("Figure: Distribution des sentiments détectés dans les avis.")
        
        # Supprimer le fichier temporaire
        try:
            os.unlink(tmp.name)
        except:
            pass
    
    # Tableau récapitulatif des avis
    doc.add_heading("Échantillon d'avis analysés", level=1)
    p = doc.add_paragraph(
        "Cette section présente un échantillon des avis analysés avec leurs scores de sentiment "
        "et autres métriques pertinentes."
    )
    
    # Déterminer les colonnes à inclure
    display_cols = ["review", "Sentiment", "Score"]
    # Ajouter quelques colonnes d'émotions principales si disponibles
    if emotion_cols:
        # Sélectionner les émotions les plus fortes en moyenne
        top_emotions = df[emotion_cols].mean().nlargest(3).index.tolist()
        display_cols.extend(top_emotions)
    
    # Limiter le nombre d'avis à 10 pour des raisons de taille
    sample_df = df.sample(min(10, len(df))) if len(df) > 10 else df
    
    # Créer le tableau
    table = doc.add_table(rows=1, cols=len(display_cols))
    table.style = "Table Grid"
    
    # En-têtes
    for i, col in enumerate(display_cols):
        header = col.replace("Emotion_", "").replace("_", " ")
        table.cell(0, i).text = header
    
    # Ajouter les lignes de données
    for _, row in sample_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(display_cols):
            if col in row:
                # Tronquer le texte des avis pour ne pas surcharger le rapport
                if col == "review" and isinstance(row[col], str):
                    text = row[col][:200] + "..." if len(row[col]) > 200 else row[col]
                    cells[i].text = text
                else:
                    value = row[col]
                    # Formatage des valeurs numériques
                    if isinstance(value, (int, float)):
                        cells[i].text = f"{value:.2f}" if isinstance(value, float) else str(value)
                    else:
                        cells[i].text = str(value)
    
    # Conclusion
    doc.add_heading("Conclusion", level=1)
    
    # Calculer quelques statistiques pour la conclusion
    conclusion_text = "L'analyse des avis révèle "
    
    if "Sentiment" in df.columns:
        sentiment_pct = sentiment_counts.get('Positif', 0) / len(df) * 100
        if sentiment_pct > 70:
            conclusion_text += f"une satisfaction globale élevée ({sentiment_pct:.1f}% d'avis positifs). "
        elif sentiment_pct > 50:
            conclusion_text += f"une satisfaction modérée ({sentiment_pct:.1f}% d'avis positifs). "
        else:
            conclusion_text += f"des défis en termes de satisfaction ({sentiment_pct:.1f}% d'avis positifs). "
    
    # Mentionner les thèmes principaux
    if keywords and len(keywords) > 0:
        conclusion_text += f"Les thèmes principaux concernent '{keywords[0]}'"
        if len(keywords) > 1:
            conclusion_text += f" et '{keywords[1]}'. "
        else:
            conclusion_text += ". "
    
    # Ajouter des détails sur les émotions si disponibles
    if emotion_cols and len(emotion_means) > 0:
        top_emotion = emotion_means.index[0].replace("Emotion_", "").replace("_", " ")
        conclusion_text += f"L'émotion dominante est '{top_emotion}', "
        conclusion_text += "ce qui suggère une expérience émotionnelle généralement "
        
        # Catégoriser les émotions en positives/négatives
        positive_emotions = ["joy", "satisfaction", "trust", "joie", "confiance", "satisfaction"]
        negative_emotions = ["anger", "fear", "sadness", "disappointment", "colère", "peur", "tristesse", "déception"]
        
        if any(emo.lower() in top_emotion.lower() for emo in positive_emotions):
            conclusion_text += "positive."
        elif any(emo.lower() in top_emotion.lower() for emo in negative_emotions):
            conclusion_text += "négative."
        else:
            conclusion_text += "nuancée."
    
    doc.add_paragraph(conclusion_text)
    
    # Recommandations (générique)
    doc.add_paragraph(
        "Les résultats suggèrent qu'une attention particulière devrait être portée aux aspects "
        "mentionnés fréquemment dans les avis, en particulier dans les commentaires négatifs. "
        "Un suivi régulier de ces analyses permettra d'évaluer l'impact des améliorations "
        "apportées au service."
    )
    
    # Sauvegarde du document en mémoire
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def generate_streamlit_report(df, keywords=None, lang="fr", include_emotions=True, include_socio=True, include_context=True):
    """
    Génère un rapport Word et le propose au téléchargement via Streamlit.
    
    Parameters:
        Identiques à export_report()
    """
    try:
        doc_io = export_report(df, keywords, lang, include_emotions, include_socio, include_context)
        
        # Créer un nom de fichier avec la date
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"Rapport_Analyse_Avis_{timestamp}.docx"
        
        # Proposer le téléchargement via Streamlit
        st.download_button(
            label="📥 Télécharger le rapport complet (DOCX)",
            data=doc_io,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.success("✅ Rapport généré avec succès ! Cliquez sur le bouton ci-dessus pour télécharger.")
        
    except Exception as e:
        st.error(f"❌ Erreur lors de la génération du rapport: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
